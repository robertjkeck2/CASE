# -*- coding: utf-8 -*-
from __future__ import absolute_import
from __future__ import division, print_function, unicode_literals
import argparse
import json
import os

from docx import Document
from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words


def parseArgs():
    parser = argparse.ArgumentParser(description='Add filename of case.')
    parser.add_argument('casefile', type=str, nargs='+', help='filename of case')
    args = parser.parse_args()
    return args.casefile

def summary(text, sentence_num):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    stemmer = Stemmer("english")

    summarizer = Summarizer(stemmer)
    summarizer.stop_words = get_stop_words("english")
    sentences = []
    for sentence in summarizer(parser.document, sentence_num):
        sentence = str(sentence).replace(',', '').split()
        for words in sentence:
            words = words.strip()
            if not words[0].isnumeric():
                start_word = words
                break
        sentence = sentence[sentence.index(start_word):]
        sentences.append(' '.join(sentence))
    return '\n'.join(sentences)

def summarize(casefile, sentence_num):
    with open('/code/results/' + casefile[:-4] + '/' + casefile[:-3] + 'json', 'r') as infile:
        case = json.loads(infile.read())
        html_file = ['<html>']
        html_file.append('<header><title>' + case['title'] + '</title><h1 style=text-align:center;font-family:Arial;>' + case['title'] + '</h1></header><body>')
        document = Document()
        html_file.append('<div style=font-family:Arial;>')
        for section in case['sections']:
            html_file.append('<h2 style=text-align:center;>' + section['title'] + '</h2>')
            document.add_heading(section['title'], 0)
            summary_text = summary(section['text'], sentence_num)
            html_file.append('<p>' + summary_text + '</p>')
            document.add_paragraph(summary_text)
        document.save('/code/results/' + casefile[:-4] + '/' + case['title'] + '.docx')
        html_file.append('</div></body></html>')
        with open('/code/results/'  + casefile[:-4] + '/' + casefile[:-3] + 'html', 'w') as html_output:
            html_output.write(''.join(html_file))

if __name__ == "__main__":
    casefile = parseArgs()
    summarize(casefile[0])
